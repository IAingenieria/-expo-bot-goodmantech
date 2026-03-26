# generar_cotizacion.py
# Goodman Tech - Generador de Cotización Profesional
# Ejecutar: python generar_cotizacion.py

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── COLORES ────────────────────────────────────────────────────────────────
AZUL_PRIMARIO   = RGBColor(0x15, 0x65, 0xC0)   # #1565C0
AZUL_OSCURO     = RGBColor(0x0D, 0x47, 0xA1)   # #0D47A1
AZUL_CLARO      = RGBColor(0xE3, 0xF2, 0xFD)   # #E3F2FD
BLANCO          = RGBColor(0xFF, 0xFF, 0xFF)
GRIS_TEXTO      = RGBColor(0x33, 0x33, 0x33)
GRIS_BORDE      = RGBColor(0xBB, 0xDE, 0xFB)   # #BBDEFB

OUTPUT_PATH = r"C:\Users\Dell\Documents\CLAUDE DESKTOP\Expo Cintermex\Cotizacion_Sistema_Telegram.docx"


# ─── HELPERS ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    """Rellena el fondo de una celda con color sólido."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    # reemplazar shd existente si lo hay
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)


def set_cell_borders(cell, color_hex="BBDEFB", sz=4):
    """Agrega bordes a una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(sz))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color_hex)
        tcBorders.append(border)
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcBorders)


def set_table_borders(table, color_hex="BBDEFB", sz=4):
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell, color_hex, sz)


def cell_para(cell, text, bold=False, font_size=10, color=None,
              alignment=WD_ALIGN_PARAGRAPH.LEFT, italic=False, space_before=0, space_after=0):
    """Limpia la celda y escribe un párrafo con formato."""
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = color
    return p


def add_run_to_para(para, text, bold=False, font_size=10, color=None, italic=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = color
    return run


def section_title(doc, text):
    """Encabezado de sección azul con texto blanco."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, AZUL_PRIMARIO)
    set_cell_borders(cell, "1565C0", 6)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"  {text}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.font.color.rgb = BLANCO
    space_after(doc, 4)


def space_after(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.font.size = Pt(pts)


def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


# ─── DOCUMENTO ──────────────────────────────────────────────────────────────

doc = Document()

# Márgenes
section = doc.sections[0]
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2)
section.right_margin  = Cm(2)

# Fuente y párrafo por defecto
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.font.color.rgb = GRIS_TEXTO

# ════════════════════════════════════════════════════════════════════════════
# ENCABEZADO — dos columnas: datos empresa | datos cotización
# ════════════════════════════════════════════════════════════════════════════
hdr_table = doc.add_table(rows=1, cols=2)
hdr_table.alignment = WD_TABLE_ALIGNMENT.LEFT
hdr_table.style = 'Table Grid'

# --- Celda izquierda: datos empresa ---
left_cell = hdr_table.cell(0, 0)
set_cell_bg(left_cell, AZUL_PRIMARIO)
set_cell_borders(left_cell, "1565C0", 6)
left_cell.width = Cm(10)

lp = left_cell.paragraphs[0]
lp.clear()
lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
lp.paragraph_format.space_before = Pt(6)
lp.paragraph_format.space_after  = Pt(2)

r = lp.add_run("Goodman Tech")
r.bold = True; r.font.size = Pt(16); r.font.name = 'Calibri'; r.font.color.rgb = BLANCO

def left_line(cell, text, bold=False, size=9):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(text)
    r.bold = bold; r.font.size = Pt(size); r.font.name = 'Calibri'; r.font.color.rgb = BLANCO

left_line(left_cell, "Soluciones Tecnológicas", bold=True, size=10)
left_line(left_cell, "Zuazua #114 Col. Centro, Monterrey N.L. CP 64000")
left_line(left_cell, "Tel: 81 2635 0902  |  info@goodmantech.com.mx")
left_line(left_cell, "RFC: GMT0000000XXX")

# Espacio final dentro de celda
p_end = left_cell.add_paragraph()
p_end.paragraph_format.space_before = Pt(4)
p_end.paragraph_format.space_after  = Pt(4)

# --- Celda derecha: número de cotización ---
right_cell = hdr_table.cell(0, 1)
set_cell_bg(right_cell, AZUL_OSCURO)
set_cell_borders(right_cell, "0D47A1", 6)
right_cell.width = Cm(7)
right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

rp = right_cell.paragraphs[0]
rp.clear()
rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
rp.paragraph_format.space_before = Pt(8)
rp.paragraph_format.space_after  = Pt(4)

rr = rp.add_run("COTIZACIÓN")
rr.bold = True; rr.font.size = Pt(18); rr.font.name = 'Calibri'; rr.font.color.rgb = BLANCO

def right_line(cell, text, bold=False, size=10):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(text)
    r.bold = bold; r.font.size = Pt(size); r.font.name = 'Calibri'; r.font.color.rgb = BLANCO

right_line(right_cell, "COT20260325-EXPO", bold=True, size=11)
right_line(right_cell, "Fecha: 25/03/2026")
right_line(right_cell, "Vigencia: 30 días")

p_end2 = right_cell.add_paragraph()
p_end2.paragraph_format.space_before = Pt(4)
p_end2.paragraph_format.space_after  = Pt(4)

space_after(doc, 8)

# ════════════════════════════════════════════════════════════════════════════
# BANNER TÍTULO DEL SISTEMA
# ════════════════════════════════════════════════════════════════════════════
banner = doc.add_table(rows=1, cols=1)
banner.alignment = WD_TABLE_ALIGNMENT.LEFT
banner.style = 'Table Grid'
bc = banner.cell(0, 0)
set_cell_bg(bc, AZUL_OSCURO)
set_cell_borders(bc, "0D47A1", 8)
bp = bc.paragraphs[0]
bp.clear()
bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
bp.paragraph_format.space_before = Pt(8)
bp.paragraph_format.space_after  = Pt(2)
br = bp.add_run("SISTEMA INTELIGENTE DE CAPTURA DE PROSPECTOS VÍA TELEGRAM")
br.bold = True; br.font.size = Pt(13); br.font.name = 'Calibri'; br.font.color.rgb = BLANCO
bp2 = bc.add_paragraph()
bp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
bp2.paragraph_format.space_before = Pt(0)
bp2.paragraph_format.space_after  = Pt(8)
br2 = bp2.add_run("Automatización Comercial con Inteligencia Artificial")
br2.italic = True; br2.font.size = Pt(10); br2.font.name = 'Calibri'; br2.font.color.rgb = BLANCO

space_after(doc, 8)

# ════════════════════════════════════════════════════════════════════════════
# SECCIÓN 1 — DATOS DEL CLIENTE
# ════════════════════════════════════════════════════════════════════════════
section_title(doc, "SECCIÓN 1 — DATOS DEL CLIENTE")

client_table = doc.add_table(rows=5, cols=2)
client_table.style = 'Table Grid'
client_table.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_borders(client_table, "BBDEFB", 4)

client_rows = [
    ("Cliente",    ""),
    ("Empresa",    "Lupront / Expo Empaques 2026"),
    ("Teléfono",   ""),
    ("Email",      ""),
    ("Evento",     "Expo Empaques 2026 — Stand 432"),
]

for i, (label, value) in enumerate(client_rows):
    row = client_table.rows[i]
    row.cells[0].width = Cm(4)
    row.cells[1].width = Cm(13)
    # encabezado de campo
    set_cell_bg(row.cells[0], AZUL_PRIMARIO if i % 2 == 0 else AZUL_OSCURO)
    cell_para(row.cells[0], label, bold=True, color=BLANCO,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=3, space_after=3)
    # valor
    if i % 2 == 0:
        set_cell_bg(row.cells[1], AZUL_CLARO)
    cell_para(row.cells[1], value, space_before=3, space_after=3)

space_after(doc, 8)

# ════════════════════════════════════════════════════════════════════════════
# SECCIÓN 2 — ¿QUÉ INCLUYE EL SISTEMA?
# ════════════════════════════════════════════════════════════════════════════
section_title(doc, "SECCIÓN 2 — ¿QUÉ INCLUYE EL SISTEMA?")

# Tabla de servicio principal
svc_table = doc.add_table(rows=2, cols=4)
svc_table.style = 'Table Grid'
svc_table.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_borders(svc_table, "BBDEFB", 4)

# Encabezados
hdrs = ["No.", "Descripción del Servicio", "Precio Unitario", "Importe"]
widths = [1.2, 10.5, 3.0, 3.0]
for j, (h, w) in enumerate(zip(hdrs, widths)):
    c = svc_table.rows[0].cells[j]
    c.width = Cm(w)
    set_cell_bg(c, AZUL_PRIMARIO)
    set_cell_borders(c, "1565C0", 6)
    cell_para(c, h, bold=True, color=BLANCO,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)

# Fila de servicio
svc_data = [
    "1",
    "Sistema Inteligente de Captura de Prospectos vía Telegram",
    "$8,000.00",
    "$8,000.00",
]
for j, (val, w) in enumerate(zip(svc_data, widths)):
    c = svc_table.rows[1].cells[j]
    c.width = Cm(w)
    set_cell_bg(c, AZUL_CLARO)
    align = WD_ALIGN_PARAGRAPH.CENTER if j != 1 else WD_ALIGN_PARAGRAPH.LEFT
    cell_para(c, val, bold=(j == 1), alignment=align, space_before=5, space_after=5)

space_after(doc, 6)

# --- Viñetas de detalle del servicio ---
bullets = [
    ("🤖 Bot de Telegram personalizado",
     " — Tu equipo captura prospectos desde el celular, sin formularios ni papeles"),
    ("🎙️ Procesamiento de audios",
     " — El vendedor graba una nota de voz y el sistema la convierte automáticamente en texto"),
    ("📷 Escaneo de tarjetas e imágenes",
     " — Toma foto de una tarjeta de presentación y el sistema extrae nombre, empresa, teléfono y correo"),
    ("📋 Base de datos en la nube",
     " — Todos los prospectos quedan organizados automáticamente con: Nombre, Empresa, Teléfono, Email, Fecha de captura, Estado del lead y Resumen de conversación"),
    ("📊 Descarga en Excel",
     " — En cualquier momento exportas toda tu lista de prospectos a Excel con un solo clic"),
    ("📧 Correo automático al prospecto",
     " — Al instante que se captura, el prospecto recibe un email personalizado con el resumen de lo conversado y los datos de contacto"),
    ("🧠 Inteligencia Artificial",
     " — El sistema analiza cada conversación y genera un resumen inteligente para dar seguimiento comercial efectivo"),
    ("⚡ Entrega en 3-5 días hábiles",
     " — Sistema listo para usar desde tu celular, sin instalaciones complicadas"),
]

for emoji_label, description in bullets:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"• {emoji_label}")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.name = 'Calibri'
    r1.font.color.rgb = AZUL_OSCURO
    r2 = p.add_run(description)
    r2.font.size = Pt(10)
    r2.font.name = 'Calibri'
    r2.font.color.rgb = GRIS_TEXTO

space_after(doc, 8)

# ════════════════════════════════════════════════════════════════════════════
# SECCIÓN 3 — ¿CÓMO FUNCIONA?
# ════════════════════════════════════════════════════════════════════════════
section_title(doc, "SECCIÓN 3 — ¿CÓMO FUNCIONA?")

steps = [
    ("1️⃣", "El vendedor abre Telegram",
     " y envía audio, foto o texto al bot — desde cualquier celular"),
    ("2️⃣", "La IA procesa la información",
     " y la registra automáticamente en la base de datos en la nube"),
    ("3️⃣", "El prospecto recibe un correo personalizado en segundos.",
     " Tú descargas tu lista en Excel cuando quieras."),
]

steps_table = doc.add_table(rows=1, cols=3)
steps_table.style = 'Table Grid'
steps_table.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_borders(steps_table, "BBDEFB", 4)

step_colors = [AZUL_PRIMARIO, AZUL_OSCURO, AZUL_PRIMARIO]
for j, (num, title, detail) in enumerate(steps):
    c = steps_table.rows[0].cells[j]
    set_cell_bg(c, step_colors[j])
    set_cell_borders(c, "1565C0", 6)

    p1 = c.paragraphs[0]
    p1.clear()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(8)
    p1.paragraph_format.space_after  = Pt(4)
    r1 = p1.add_run(num)
    r1.font.size = Pt(20); r1.font.name = 'Calibri'; r1.font.color.rgb = BLANCO

    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(2)
    r2 = p2.add_run(title)
    r2.bold = True; r2.font.size = Pt(10); r2.font.name = 'Calibri'; r2.font.color.rgb = BLANCO

    p3 = c.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after  = Pt(8)
    r3 = p3.add_run(detail)
    r3.font.size = Pt(9); r3.font.name = 'Calibri'; r3.font.color.rgb = BLANCO; r3.italic = True

space_after(doc, 8)

# ════════════════════════════════════════════════════════════════════════════
# SECCIÓN 4 — CAMPOS QUE SE CAPTURAN AUTOMÁTICAMENTE
# ════════════════════════════════════════════════════════════════════════════
section_title(doc, "SECCIÓN 4 — CAMPOS QUE SE CAPTURAN AUTOMÁTICAMENTE")

fields = [
    ("Nombre",                  "Nombre completo del prospecto"),
    ("Email",                   "Correo electrónico de contacto"),
    ("Empresa",                 "Empresa u organización del prospecto"),
    ("Teléfono",                "Número de contacto directo"),
    ("Fecha de Captura",        "Fecha y hora exacta del registro"),
    ("Fuente",                  "Canal de origen: Expo, Web, Referido, etc."),
    ("Estado del Lead",         "Nuevo / En conversación / Calificado / Cerrado"),
    ("Temperatura",             "Frío / Tibio / Caliente — evaluado por IA"),
    ("Puntuación",              "Score de 0 a 100 generado automáticamente por IA"),
    ("Resumen de conversación", "Texto generado automáticamente por Inteligencia Artificial"),
    ("Última interacción",      "Fecha del último contacto registrado"),
]

fields_table = doc.add_table(rows=len(fields) + 1, cols=2)
fields_table.style = 'Table Grid'
fields_table.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_borders(fields_table, "BBDEFB", 4)

# Encabezado
for j, h in enumerate(["Campo", "Descripción"]):
    c = fields_table.rows[0].cells[j]
    set_cell_bg(c, AZUL_PRIMARIO)
    set_cell_borders(c, "1565C0", 6)
    cell_para(c, h, bold=True, color=BLANCO,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)

fields_table.rows[0].cells[0].width = Cm(5)
fields_table.rows[0].cells[1].width = Cm(12.7)

for i, (field, desc) in enumerate(fields):
    row = fields_table.rows[i + 1]
    row.cells[0].width = Cm(5)
    row.cells[1].width = Cm(12.7)
    bg = AZUL_CLARO if i % 2 == 0 else BLANCO
    set_cell_bg(row.cells[0], AZUL_OSCURO if i % 2 == 0 else AZUL_PRIMARIO)
    set_cell_bg(row.cells[1], bg)
    cell_para(row.cells[0], field, bold=True, color=BLANCO,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=3, space_after=3)
    cell_para(row.cells[1], desc, space_before=3, space_after=3)

space_after(doc, 8)

# ════════════════════════════════════════════════════════════════════════════
# SECCIÓN 5 — INVERSIÓN
# ════════════════════════════════════════════════════════════════════════════
section_title(doc, "SECCIÓN 5 — INVERSIÓN")

inv_table = doc.add_table(rows=3, cols=2)
inv_table.style = 'Table Grid'
inv_table.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_borders(inv_table, "BBDEFB", 4)

inv_data = [
    ("Subtotal",    "$8,000.00 MXN",  False, BLANCO,       AZUL_CLARO),
    ("IVA (16%)",   "$1,280.00 MXN",  False, BLANCO,       BLANCO),
    ("TOTAL",       "$9,280.00 MXN",  True,  BLANCO,       AZUL_OSCURO),
]

for i, (label, amount, bold, txt_color, bg_right) in enumerate(inv_data):
    row = inv_table.rows[i]
    row.cells[0].width = Cm(10)
    row.cells[1].width = Cm(7.7)
    set_cell_bg(row.cells[0], AZUL_PRIMARIO if i < 2 else AZUL_OSCURO)
    set_cell_bg(row.cells[1], bg_right)
    cell_para(row.cells[0], label, bold=bold, color=BLANCO,
              alignment=WD_ALIGN_PARAGRAPH.LEFT, font_size=11 if bold else 10,
              space_before=5, space_after=5)
    cell_para(row.cells[1], amount, bold=bold, color=BLANCO if bold else GRIS_TEXTO,
              alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_size=12 if bold else 10,
              space_before=5, space_after=5)

space_after(doc, 8)

# ════════════════════════════════════════════════════════════════════════════
# SECCIÓN 6 — CONDICIONES DE PAGO / DATOS BANCARIOS
# ════════════════════════════════════════════════════════════════════════════
section_title(doc, "SECCIÓN 6 — CONDICIONES DE PAGO Y DATOS BANCARIOS")

pay_table = doc.add_table(rows=1, cols=2)
pay_table.style = 'Table Grid'
pay_table.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_borders(pay_table, "BBDEFB", 4)

# --- Condiciones de pago (izquierda) ---
left = pay_table.rows[0].cells[0]
left.width = Cm(8.5)
set_cell_bg(left, AZUL_CLARO)
set_cell_borders(left, "BBDEFB", 4)

lp0 = left.paragraphs[0]
lp0.clear()
lp0.alignment = WD_ALIGN_PARAGRAPH.LEFT
lp0.paragraph_format.space_before = Pt(6)
lp0.paragraph_format.space_after  = Pt(4)
r0 = lp0.add_run("CONDICIONES DE PAGO")
r0.bold = True; r0.font.size = Pt(11); r0.font.name = 'Calibri'; r0.font.color.rgb = AZUL_OSCURO

pay_bullets = [
    "• 50% de anticipo para iniciar: $4,640.00 MXN",
    "• 50% restante al finalizar: $4,640.00 MXN",
    "• Precios expresados en Pesos Mexicanos (MXN)",
    "• Entrega en 3 a 5 días hábiles",
    "• Soporte técnico incluido por 30 días",
]
for line in pay_bullets:
    pp = left.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pp.paragraph_format.space_before = Pt(2)
    pp.paragraph_format.space_after  = Pt(2)
    rr = pp.add_run(line)
    rr.font.size = Pt(10); rr.font.name = 'Calibri'; rr.font.color.rgb = GRIS_TEXTO

p_pad = left.add_paragraph()
p_pad.paragraph_format.space_before = Pt(4)
p_pad.paragraph_format.space_after  = Pt(4)

# --- Datos bancarios (derecha) ---
right = pay_table.rows[0].cells[1]
right.width = Cm(9.2)
set_cell_bg(right, AZUL_OSCURO)
set_cell_borders(right, "0D47A1", 6)

rp0 = right.paragraphs[0]
rp0.clear()
rp0.alignment = WD_ALIGN_PARAGRAPH.LEFT
rp0.paragraph_format.space_before = Pt(6)
rp0.paragraph_format.space_after  = Pt(4)
rr0 = rp0.add_run("DATOS BANCARIOS")
rr0.bold = True; rr0.font.size = Pt(11); rr0.font.name = 'Calibri'; rr0.font.color.rgb = BLANCO

bank_lines = [
    ("Beneficiario: ", "Guadalupe Salinas Baños"),
    ("Banco: ",        "MERCADO PAGO"),
    ("CLABE: ",        "7229 6901 0455 5450 66"),
]
for label, value in bank_lines:
    bp_line = right.add_paragraph()
    bp_line.alignment = WD_ALIGN_PARAGRAPH.LEFT
    bp_line.paragraph_format.space_before = Pt(3)
    bp_line.paragraph_format.space_after  = Pt(3)
    rl = bp_line.add_run(label)
    rl.bold = True; rl.font.size = Pt(10); rl.font.name = 'Calibri'; rl.font.color.rgb = AZUL_CLARO
    rv = bp_line.add_run(value)
    rv.font.size = Pt(10); rv.font.name = 'Calibri'; rv.font.color.rgb = BLANCO

# CLABE resaltada en bloque separado
clabe_block = right.add_paragraph()
clabe_block.alignment = WD_ALIGN_PARAGRAPH.CENTER
clabe_block.paragraph_format.space_before = Pt(6)
clabe_block.paragraph_format.space_after  = Pt(6)
rc = clabe_block.add_run("7229 6901 0455 5450 66")
rc.bold = True; rc.font.size = Pt(13); rc.font.name = 'Calibri'; rc.font.color.rgb = BLANCO

space_after(doc, 8)

# ════════════════════════════════════════════════════════════════════════════
# SECCIÓN 7 — TÉRMINOS Y CONDICIONES
# ════════════════════════════════════════════════════════════════════════════
section_title(doc, "SECCIÓN 7 — TÉRMINOS Y CONDICIONES")

terms = [
    "El sistema se entrega configurado con la información de tu empresa.",
    "El cliente es responsable de las cuentas de Telegram, Gmail y los servicios de IA (costo aproximado $10–30 USD/mes según volumen de uso).",
    "El anticipo del 50% confirma el pedido; el 50% restante se paga al momento de la entrega.",
    "El precio incluye una sola instalación y configuración inicial.",
    "No incluye mantenimiento mensual (disponible por separado bajo cotización adicional).",
    "La vigencia de esta cotización es de 30 días naturales a partir de la fecha de emisión.",
    "Precios expresados en Pesos Mexicanos (MXN) más IVA.",
    "Entrega vía instalación remota; capacitación por videollamada incluida (30 minutos).",
]

for i, term in enumerate(terms):
    tp = doc.add_paragraph()
    tp.paragraph_format.left_indent = Cm(0.5)
    tp.paragraph_format.space_before = Pt(2)
    tp.paragraph_format.space_after  = Pt(2)
    rn = tp.add_run(f"{i+1}. ")
    rn.bold = True; rn.font.size = Pt(10); rn.font.name = 'Calibri'; rn.font.color.rgb = AZUL_OSCURO
    rt = tp.add_run(term)
    rt.font.size = Pt(10); rt.font.name = 'Calibri'; rt.font.color.rgb = GRIS_TEXTO

space_after(doc, 8)

# ════════════════════════════════════════════════════════════════════════════
# PIE DE PÁGINA
# ════════════════════════════════════════════════════════════════════════════
footer_table = doc.add_table(rows=2, cols=1)
footer_table.style = 'Table Grid'
footer_table.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_borders(footer_table, "1565C0", 6)

# Fila 1 — vigencia
vc = footer_table.rows[0].cells[0]
set_cell_bg(vc, AZUL_CLARO)
cell_para(vc,
          "Esta cotización tiene una vigencia de 30 días a partir de la fecha de emisión.",
          bold=True, color=AZUL_OSCURO,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=5, space_after=5)

# Fila 2 — datos empresa
fc = footer_table.rows[1].cells[0]
set_cell_bg(fc, AZUL_PRIMARIO)
cell_para(fc,
          "Goodman Tech  |  Zuazua #114 Col. Centro, Monterrey N.L. CP 64000  |  Tel: 81 2635 0902  |  info@goodmantech.com.mx",
          bold=False, font_size=9, color=BLANCO,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)

# ════════════════════════════════════════════════════════════════════════════
# GUARDAR
# ════════════════════════════════════════════════════════════════════════════
doc.save(OUTPUT_PATH)
print("Cotizacion generada exitosamente en:")
print(f"   {OUTPUT_PATH}")
